import json
from typing import List, Optional, Tuple, Union
from pydantic import BaseModel, Field
from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime
from zoneinfo import ZoneInfo
from pathlib import Path

import torch
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

TIMEZONE = "Europe/Moscow"
OUT_DIR = Path("out"); OUT_DIR.mkdir(exist_ok=True)
DOC_TITLE = "Документ установки трансфертного ценообразования\n\nмежду ООО Фибратек и ООО МИНЕРАЛ КОМПОЗИТ"
ECONOMIST_LINE = "Утверждено экономистом ООО МИНЕРАЛ КОМПОЗИТ"

DISCOUNT_PCT = 12.0

COND_ADJUST_PCT = {
    "Кондиция": 0.0,
    "Некондиция 1 сорт": -2.0,
    "Некондиция": -5.0,
}

COND_ORDER = ["Кондиция", "Некондиция 1 сорт", "Некондиция"]

MODEL_ID = "meta-llama/Meta-Llama-3.1-8B-Instruct"

class Item(BaseModel):
    nomenclature: str = Field(..., description="Например: 'ЛПН 3000х1200х10 ТУ (Серый, не грунт)'")
    sale_price_rub: float = Field(..., description="Цена реализации (число)")
    condition: Optional[Union[str, List[str]]] = Field(
        None,
        description="'Кондиция'|'Некондиция'|'Некондиция 1 сорт'|или список таких строк|null"
    )

class Extraction(BaseModel):
    items: List[Item]

bnb = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_use_double_quant=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype=torch.bfloat16,
)

tok = AutoTokenizer.from_pretrained(MODEL_ID, use_fast=True)
model = AutoModelForCausalLM.from_pretrained(
    MODEL_ID,
    device_map="auto",
    quantization_config=bnb,
    dtype=torch.bfloat16,  # вместо torch_dtype
)

SYSTEM = (
    "Ты извлекаешь позиции из русских деловых сообщений. "
    "Верни СТРОГО JSON по схеме: "
    '{"items":[{"nomenclature":"строка","sale_price_rub":число,'
    '"condition":"Кондиция|Некондиция|Некондиция 1 сорт|[строки]|null"}]}. '
    "Если несколько НОМЕНКЛАТУР — создай несколько объектов в items. "
    "Сопоставляй цену каждой номенклатуре по принципу БЛИЖАЙШЕГО ЧИСЛА-ЦЕНЫ, встречающегося рядом справа, "
    "после двоеточия или в той же фразе/строке. "
    "Если для ОДНОЙ номенклатуры указаны НЕСКОЛЬКО СОСТОЯНИЙ (например Кондиция, Некондиция, Некондиция 1 сорт): "
    "— и есть НЕСКОЛЬКО ЦЕН рядом: создай ОТДЕЛЬНЫЕ объекты в items с соответствующими ценами; "
    "при сопоставлении более ВЫСОКАЯ цена относится к 'Кондиция', более НИЗКАЯ — к 'Некондиция'. "
    "— если указана ОДНА цена на все состояния: верни одно items-значение с 'condition' как МАССИВ строк, "
    "например: \"condition\":[\"Кондиция\",\"Некондиция\"] (цены будут одинаковые до дальнейшей пост-обработки). "
    "Если цена встречается несколько раз — возьми ПОСЛЕДНЮЮ, находящуюся ближе всего к номенклатуре. "
    "Размеры могут быть вида 3000х1200х10 (русская 'х'). "
    "Стандарты: ГОСТ или ТУ. "
    "Если в сообщении встречаются дополнительные характеристики в скобках (например 'Серый, не грунт'), "
    "включай их прямо в поле nomenclature как часть строки. "
    "Если условие кондиции/некондиции явно не указано — заполни condition значением 'Кондиция'. "
    "Никаких комментариев — ТОЛЬКО JSON."
)

def ask_json(text: str, max_new_tokens: int = 512) -> Extraction:
    messages = [
        {"role": "system", "content": SYSTEM},
        {"role": "user", "content": text},
    ]
    prompt = tok.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    inputs = tok([prompt], return_tensors="pt").to(model.device)

    with torch.inference_mode():
        output_ids = model.generate(
            **inputs,
            max_new_tokens=max_new_tokens,
            do_sample=False,
            temperature=0.0,
            eos_token_id=tok.eos_token_id,
        )

    gen = tok.decode(output_ids[0][inputs["input_ids"].shape[1]:], skip_special_tokens=True)
    s, e = gen.find("{"), gen.rfind("}")
    payload = gen[s:e+1] if s != -1 and e != -1 else "{}"
    data = json.loads(payload)
    return Extraction(**data)

# ===== Утилиты DOCX =====
def fmt_rub(amount: Decimal) -> str:
    q = amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    s = f"{q:.2f}".replace(",", "X").replace(".", ",").replace("X",".")
    ip, fp = s.split(",")
    ip = " ".join([ip[max(i-3,0):i] for i in range(len(ip),0,-3)][::-1])
    return f"{ip},{fp}"

def style_table(table):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(12)

def safe_filename(name: str) -> str:
    return "".join(c if c not in '\\/:*?"<>|' else "_" for c in name)

def now_date_str() -> str:
    try:
        now = datetime.now(ZoneInfo(TIMEZONE))
    except Exception:
        now = datetime.now()
    return now.strftime("%d.%m.%Y")

def normalize_condition(cond: Optional[str]) -> str:
    if not cond or not str(cond).strip():
        return "Кондиция"
    t = str(cond).strip().lower()
    if "некондиц" in t and "1" in t:
        return "Некондиция 1 сорт"
    if "некондиц" in t:
        return "Некондиция"
    if "кондиц" in t:
        return "Кондиция"
    return "Кондиция"

def adjust_sale_price_by_condition(base_price: Decimal, cond: str) -> Decimal:
    adj_pct = Decimal(str(COND_ADJUST_PCT.get(cond, 0.0)))
    return base_price * (Decimal("1.0") + adj_pct / Decimal("100"))

# ===== Документы =====
def make_docx_single(nomenclature: str, condition: Optional[str], transfer_price: Decimal) -> Path:
    doc = Document()

    # Заголовок по центру
    p = doc.add_paragraph(DOC_TITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.size = Pt(12)

    doc.add_paragraph("")
    p2 = doc.add_paragraph("Прошу установить новую трансфертную цену на номенклатуру")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p2.runs: r.font.size = Pt(12)
    doc.add_paragraph("")

    # Таблица
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Номенклатура"
    table.rows[0].cells[1].text = "Характеристика"
    table.rows[0].cells[2].text = "Трансфертная цена, руб (вкл НДС)"
    for j in range(3):
        hdr_p = table.rows[0].cells[j].paragraphs[0]
        hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in hdr_p.runs:
            r.font.bold = True
            r.font.size = Pt(12)

    table.rows[1].cells[0].text = nomenclature
    table.rows[1].cells[1].text = normalize_condition(condition)
    table.rows[1].cells[2].text = fmt_rub(transfer_price)
    style_table(table)

    doc.add_paragraph(""); doc.add_paragraph("")

    p3 = doc.add_paragraph(ECONOMIST_LINE)
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p3.runs: r.font.size = Pt(12)

    today = now_date_str()
    p4 = doc.add_paragraph(f"Дата\t\t\t\t\t\t\t\t{today}")
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p4.runs: r.font.size = Pt(12)

    out_name = safe_filename(f"{nomenclature} {today}.docx")
    out = OUT_DIR / out_name
    doc.save(out)
    return out

def make_docx_multi(rows: List[Tuple[str, Optional[str], Decimal]]) -> Path:
    doc = Document()

    p = doc.add_paragraph(DOC_TITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.size = Pt(12)

    doc.add_paragraph("")
    p2 = doc.add_paragraph("Прошу установить новую трансфертную цену на номенклатуры")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p2.runs: r.font.size = Pt(12)
    doc.add_paragraph("")

    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.rows[0].cells[0].text = "Номенклатура"
    table.rows[0].cells[1].text = "Характеристика"
    table.rows[0].cells[2].text = "Трансфертная цена, руб (вкл НДС)"
    for j in range(3):
        hdr_p = table.rows[0].cells[j].paragraphs[0]
        hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in hdr_p.runs:
            r.font.bold = True
            r.font.size = Pt(12)

    for i, (nomenclature, condition, transfer_price) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = nomenclature
        table.rows[i].cells[1].text = normalize_condition(condition)
        table.rows[i].cells[2].text = fmt_rub(transfer_price)

    style_table(table)

    doc.add_paragraph(""); doc.add_paragraph("")

    p3 = doc.add_paragraph(ECONOMIST_LINE)
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p3.runs: r.font.size = Pt(12)

    today = now_date_str()
    p4 = doc.add_paragraph(f"Дата\t\t\t\t\t\t\t\t{today}")
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p4.runs: r.font.size = Pt(12)

    out = OUT_DIR / f"Трансфертные цены {today}.docx"
    doc.save(out)
    return out

def expand_rows_for_item(it: Item, discount_pct: float) -> List[Tuple[str, Optional[str], Decimal]]:
    """
    Превращает один Item (с condition-строкой или списком) в одну или несколько строк для таблицы.
    Если condition — список (и была одна базовая цена), применяем корректировки COND_ADJUST_PCT,
    чтобы 'Кондиция' была дороже.
    """
    base_price = Decimal(str(it.sale_price_rub))
    rows: List[Tuple[str, Optional[str], Decimal]] = []

    if it.condition is None or (isinstance(it.condition, str) and not it.condition.strip()):
        conds = ["Кондиция"]
    elif isinstance(it.condition, str):
        conds = [it.condition]
    else:
        conds = it.condition if it.condition else ["Кондиция"]

    norm_conds = [normalize_condition(c) for c in conds]
    norm_conds = sorted(norm_conds, key=lambda c: COND_ORDER.index(c) if c in COND_ORDER else 999)

    multiple_conditions = len(norm_conds) > 1

    for cond in norm_conds:
        effective_sale_price = base_price
        if multiple_conditions:
            effective_sale_price = adjust_sale_price_by_condition(base_price, cond)
        k = Decimal(str(1 - discount_pct/100.0))
        transfer = (effective_sale_price * k)
        rows.append((it.nomenclature, cond, transfer))

    return rows

def run_agent(raw_text: str, discount_pct: Optional[float] = None,
              combine: Optional[bool] = None, max_new_tokens: int = 512) -> List[str]:
    if discount_pct is None:
        discount_pct = DISCOUNT_PCT

    result = ask_json(raw_text, max_new_tokens=max_new_tokens)
    created: List[str] = []
    rows: List[Tuple[str, Optional[str], Decimal]] = []

    if combine is None:
        combine = len(result.items) > 1

    for it in result.items:
        for name, cond, transfer in expand_rows_for_item(it, discount_pct):
            if combine:
                rows.append((name, cond, transfer))
            else:
                path = make_docx_single(name, cond, transfer)
                print("Создан файл:", path)
                created.append(str(path))

    if combine and rows:
        path = make_docx_multi(rows)
        print("Создан файл:", path)
        created.append(str(path))

    return created

if __name__ == "__main__":
    import argparse
    p = argparse.ArgumentParser()
    p.add_argument("--text", type=str, required=True, help="вставь текст сообщений (можно с несколькими позициями)")
    p.add_argument("--discount", type=float, default=None, help="скидка в процентах (если не указана — берётся DISCOUNT_PCT)")
    p.add_argument("--max_new_tokens", type=int, default=512, help="лимит генерации для LLM")
    p.add_argument("--split", action="store_true", help="каждую позицию в отдельный документ (по умолчанию — объединяем при >1)")
    args = p.parse_args()

    combine_flag = not args.split

    run_agent(
        args.text,
        discount_pct=args.discount,
        combine=combine_flag,
        max_new_tokens=args.max_new_tokens
    )
